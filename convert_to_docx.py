# -*- coding: utf-8 -*-
import glob
import os

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


class Src2Docx():

    def __init__(self, template, name_doc, name_num_dec):
        self.doc = Document(template)
        section = self.doc.sections[0]
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

        self.add_name_and_dec_num(template, name_doc, name_num_dec)

    def add_name_and_dec_num(self, template, name_doc, name_num_dec):
        tmp_name_doc = 'ИМЯ ДОКУМЕНТА'
        tmp_name_dec = 'ДЕЦ.НОМЕР'
        # проходимся по таблицам
        for i, table in enumerate(self.doc.tables):
            # проходимся по строкам таблицы `i`
            for j, row in enumerate(table.rows):
                # проходимся по ячейкам таблицы `i` и строки `j`
                for cell in row.cells:
                    # добавляем значение ячейки в соответствующий
                    # список, созданного словаря под данные таблиц
                    txt_val = cell.text
                    if len(txt_val):
                        txt_val = txt_val.replace(tmp_name_dec, name_num_dec)
                        cell.text = txt_val

                    paragraphs = cell.paragraphs

                    for paragraph in paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(12)
        for para in self.doc.paragraphs:
            if para.text == tmp_name_doc:
                para.text = name_doc
            elif tmp_name_dec in para.text:
                txt_val = para.text
                txt_val = txt_val.replace(tmp_name_dec, name_num_dec)
                para.text = txt_val

    def add_heading(self, file_bn):
        style = self.doc.styles['Heading 1']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        rFonts = style.element.rPr.rFonts
        rFonts.set(qn("w:asciiTheme"), "Times New Roman")
        rFonts.set(qn("w:hAnsiTheme"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        head = self.doc.add_heading(file_bn, 1)
        h_fmt = head.paragraph_format
        h_fmt.line_spacing = 1.5
        h_fmt.space_before = Pt(0)
        h_fmt.space_after = Pt(0)
        h_fmt.first_line_indent = Cm(1.5)

    def add_paragraph(self, text):
        style = self.doc.styles['Normal']
        style.font.size = Pt(14)
        style.font.name = 'Times New Roman'
        p = self.doc.add_paragraph(text)
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        p_fmt = p.paragraph_format
        p_fmt.line_spacing = 1.5
        p_fmt.space_before = Pt(0)
        p_fmt.space_after = Pt(0)
        p_fmt.first_line_indent = Cm(1.5)

    def add_files(self, files):
        for i in files:
            name_f = os.path.basename(i)
            f = open(i,'r', encoding="utf-8")#, errors='ignore')
            try:
                text_str = f.read()
                f.close()
                self.add_heading(name_f)
                self.add_paragraph(text_str)
            except Exception:
                print("Error file utf8:", i)

    def add_koll(self, texth):
        header = self.doc.sections[0].header.paragraphs[0]
        header.text = texth
        header.style.font.size = Pt(14)
        header.style.font.name = 'Times New Roman'
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_table_lri(self):
        path_lri = ".\\auxiliary\\template_lri.docx"
        doc_lri = Document(path_lri)

        for table in doc_lri.tables:
            p = self.doc.add_paragraph()
            p._p.addnext(table._tbl)
            if len(p.text) == 0:
                p = p._element
                p.getparent().remove(p)
                p._p = p._element = None

    def save_docx(self, path):
        self.doc.save(path)
